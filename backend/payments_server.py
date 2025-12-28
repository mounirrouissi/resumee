"""
FastAPI payment server for Stripe Checkout credit purchasing.
Handles checkout session creation, webhook processing, and credit-gated document generation.
"""
import os
import logging
from typing import Optional
from pathlib import Path

from fastapi import FastAPI, HTTPException, Request, Form, File, UploadFile
from fastapi.responses import JSONResponse, FileResponse
from pydantic import BaseModel, Field, validator
from dotenv import load_dotenv
import stripe
from stripe import StripeError

from backend import db
from backend import stripe_utils
from docx import Document
from docx.shared import Pt
from docx.text.paragraph import Paragraph

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Initialize FastAPI app
app = FastAPI(title="Resume Credits Payment API")

# Environment configuration
BASE_URL = os.getenv("BASE_URL", "https://resumee-nhrs.onrender.com")
STRIPE_SECRET_KEY = os.getenv("STRIPE_SECRET_KEY")
STRIPE_WEBHOOK_SECRET = os.getenv("STRIPE_WEBHOOK_SECRET")

if not STRIPE_SECRET_KEY:
    logger.warning("STRIPE_SECRET_KEY not set - payment features will not work")


# ==================== Request Models ====================

class CreateCheckoutRequest(BaseModel):
    """Request model for creating a Stripe Checkout session."""
    session_id: str = Field(..., description="Application session ID (UUID format)")
    pack: str = Field(..., description="Credit pack: '1', '5', or '12'")
    
    @validator('session_id')
    def validate_session_id(cls, v):
        """Ensure session_id looks like a UUID."""
        if len(v) < 10 or not all(c.isalnum() or c == '-' for c in v):
            raise ValueError("session_id must be UUID-like format")
        return v
    
    @validator('pack')
    def validate_pack(cls, v):
        """Ensure pack is a valid option."""
        if v not in ["1", "5", "12"]:
            raise ValueError("pack must be '1', '5', or '12'")
        return v


# ==================== Endpoints ====================

@app.get("/")
async def root():
    """Health check endpoint."""
    return {
        "service": "Resume Credits Payment API",
        "status": "running",
        "stripe_configured": bool(STRIPE_SECRET_KEY)
    }


@app.post("/create-checkout")
async def create_checkout(request: CreateCheckoutRequest):
    """
    Create a Stripe Checkout session for purchasing credits.
    
    Returns checkout URL for the user to complete payment.
    """
    if not STRIPE_SECRET_KEY:
        raise HTTPException(
            status_code=503,
            detail="Payment service not configured"
        )
    
    try:
        # Verify session exists (or create it)
        session_data = db.get_session(request.session_id)
        if not session_data:
            # Create session with 0 credits initially
            db.create_session(request.session_id, improved_text="", credits=0)
            logger.info(f"Created new session for checkout: {request.session_id}")
        
        # Generate success and cancel URLs
        success_url = f"{BASE_URL}/success?session_id={request.session_id}"
        cancel_url = f"{BASE_URL}/cancel?session_id={request.session_id}"
        
        # Create Stripe Checkout session
        result = stripe_utils.create_checkout_session(
            session_id=request.session_id,
            pack=request.pack,
            success_url=success_url,
            cancel_url=cancel_url
        )
        
        return {
            "checkout_url": result["checkout_url"],
            "session_id": request.session_id,
            "pack": request.pack
        }
    
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except StripeError as e:
        logger.error(f"Stripe error: {e}")
        raise HTTPException(status_code=500, detail="Payment service error")


@app.post("/webhook")
async def stripe_webhook(request: Request):
    """
    Handle Stripe webhook events.
    
    Processes checkout.session.completed events to add credits to user accounts.
    """
    if not STRIPE_WEBHOOK_SECRET:
        logger.error("STRIPE_WEBHOOK_SECRET not configured")
        raise HTTPException(status_code=503, detail="Webhook not configured")
    
    # Get raw body and signature
    payload = await request.body()
    signature = request.headers.get("Stripe-Signature")
    
    if not signature:
        raise HTTPException(status_code=400, detail="Missing Stripe-Signature header")
    
    try:
        # Verify webhook signature
        event = stripe_utils.verify_webhook_signature(
            payload=payload,
            signature=signature,
            webhook_secret=STRIPE_WEBHOOK_SECRET
        )
        
        # Handle checkout.session.completed event
        if event["type"] == "checkout.session.completed":
            session = event["data"]["object"]
            metadata = session.get("metadata", {})
            
            app_session_id = metadata.get("app_session_id")
            credits_str = metadata.get("credits")
            
            if not app_session_id or not credits_str:
                logger.error(f"Missing metadata in webhook: {metadata}")
                return JSONResponse(
                    status_code=400,
                    content={"error": "invalid_metadata"}
                )
            
            try:
                credits = int(credits_str)
            except ValueError:
                logger.error(f"Invalid credits value: {credits_str}")
                return JSONResponse(
                    status_code=400,
                    content={"error": "invalid_credits"}
                )
            
            # Add credits to session (creates session if doesn't exist)
            db.add_credits(app_session_id, credits)
            
            logger.info(
                f"Webhook processed: Added {credits} credits to session {app_session_id}"
            )
            
            return JSONResponse(
                status_code=200,
                content={
                    "status": "success",
                    "session_id": app_session_id,
                    "credits_added": credits
                }
            )
        
        # Acknowledge other event types
        logger.info(f"Received webhook event: {event['type']}")
        return JSONResponse(status_code=200, content={"status": "received"})
    
    except StripeError as e:
        logger.error(f"Invalid webhook signature: {e}")
        raise HTTPException(status_code=400, detail="Invalid signature")
    except Exception as e:
        logger.error(f"Webhook processing error: {e}")
        raise HTTPException(status_code=500, detail="Webhook processing failed")


@app.post("/generate")
async def generate_document(
    session_id: str = Form(...),
    filename: str = Form("improved_resume.docx")
):
    """
    Generate and download improved resume document.
    
    Requires credits. If no credits available, returns 402 with checkout URL.
    If credits available, deducts one credit and returns the document.
    """
    # Get session data
    session_data = db.get_session(session_id)
    
    if not session_data:
        raise HTTPException(
            status_code=404,
            detail={
                "error": "session_not_found",
                "message": "Session not found. Please upload a resume first."
            }
        )
    
    _, improved_text, credits = session_data
    
    # Check if credits available
    if credits <= 0:
        # Generate checkout URL for default pack (5 credits)
        checkout_url = stripe_utils.get_checkout_url_for_session(
            session_id=session_id,
            pack="5",
            base_url=BASE_URL
        )
        
        return JSONResponse(
            status_code=402,
            content={
                "error": "no_credits",
                "message": "No credits available. Purchase credits to download.",
                "checkout_url": checkout_url,
                "session_id": session_id,
                "credits": 0
            }
        )
    
    # Deduct credit atomically
    success = db.deduct_credit(session_id)
    
    if not success:
        # Race condition: credits depleted between check and deduct
        checkout_url = stripe_utils.get_checkout_url_for_session(
            session_id=session_id,
            pack="5",
            base_url=BASE_URL
        )
        
        return JSONResponse(
            status_code=402,
            content={
                "error": "no_credits",
                "message": "No credits available. Purchase credits to download.",
                "checkout_url": checkout_url,
                "session_id": session_id,
                "credits": 0
            }
        )
    
    # Generate DOCX file
    # TODO: Integrate with existing OCR/generation logic here
    # For now, create a simple document from improved_text
    
    output_path = Path("backend/outputs") / filename
    output_path.parent.mkdir(exist_ok=True)
    
    doc = Document()
    
    if improved_text:
        # Add improved text to document
        for paragraph_text in improved_text.split('\n\n'):
            if paragraph_text.strip():
                doc.add_paragraph(paragraph_text.strip())
    else:
        # Placeholder if no improved text
        doc.add_paragraph("Your improved resume will appear here.")
        doc.add_paragraph("Please upload and process your resume first.")
    
    doc.save(str(output_path))
    
    logger.info(f"Generated document for session {session_id}, credits remaining: {credits - 1}")
    
    # Return file for download
    return FileResponse(
        path=str(output_path),
        filename=filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


@app.post("/generate-watermarked")
async def generate_watermarked_document(
    session_id: str = Form(...),
    filename: str = Form("resume_sample.docx")
):
    """
    Generate a watermarked sample document (no credits required).
    
    Useful for preview purposes. Every paragraph has a watermark.
    """
    # Get session data
    session_data = db.get_session(session_id)
    
    if not session_data:
        raise HTTPException(
            status_code=404,
            detail={
                "error": "session_not_found",
                "message": "Session not found. Please upload a resume first."
            }
        )
    
    _, improved_text, _ = session_data
    
    # Generate watermarked DOCX file
    output_path = Path("backend/outputs") / filename
    output_path.parent.mkdir(exist_ok=True)
    
    doc = Document()
    
    watermark_text = "-- WATERMARKED SAMPLE — PURCHASE TO DOWNLOAD ORIGINAL --"
    
    if improved_text:
        # Add improved text with watermarks
        for paragraph_text in improved_text.split('\n\n'):
            if paragraph_text.strip():
                doc.add_paragraph(paragraph_text.strip())
                # Add watermark after each paragraph
                doc.add_paragraph(watermark_text)
    else:
        # Placeholder if no improved text
        doc.add_paragraph("Your improved resume will appear here.")
        doc.add_paragraph(watermark_text)
    
    doc.save(str(output_path))
    
    logger.info(f"Generated watermarked document for session {session_id}")
    
    # Return file for download
    return FileResponse(
        path=str(output_path),
        filename=filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


@app.get("/session/{session_id}")
async def get_session_info(session_id: str):
    """Get session information including credit balance."""
    session_data = db.get_session(session_id)
    
    if not session_data:
        raise HTTPException(
            status_code=404,
            detail={"error": "session_not_found"}
        )
    
    _, improved_text, credits = session_data
    
    return {
        "session_id": session_id,
        "credits": credits,
        "has_improved_text": bool(improved_text)
    }


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000, reload=True)
